#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""Generate patent application document as DOCX based on DeepDebateCoder paper."""

import json
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# --- Paths ---
SKILL_DIR = Path.cwd()
RUN_ID = "KV61EIpzVvuNGPGZ"
OUTPUT_DIR = SKILL_DIR / "output" / "runs" / RUN_ID
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# --- Helper functions ---
def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elem)

def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return heading

def add_paragraph_styled(doc, text, bold=False, font_size=12, alignment=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = Pt(22)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

# --- Create document ---
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# --- Title page ---
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('基于非对称多智能体辩论的\n可靠可维护代码生成方法及系统')
title_run.font.size = Pt(22)
title_run.bold = True
title_run.font.name = '黑体'
title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
title_p.paragraph_format.space_after = Pt(30)

# Add empty lines
doc.add_paragraph()
doc.add_paragraph()

# Application info
info_text = [
    "申请人：太原理工大学",
    "发明人：王云鹏、韩阳志、高佳英、张文瑞、王剑锋、张伟",
    "申请日：2025年",
    "摘要：本发明公开了一种基于非对称多智能体辩论的可靠可维护代码生成方法及系统。"
]
for txt in info_text:
    add_paragraph_styled(doc, txt, font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ========== 摘要 ==========
add_heading_styled(doc, '摘要', level=1)
abstract_text = (
    "本发明公开了一种基于非对称多智能体辩论的可靠可维护代码生成方法及系统，属于人工智能与代码生成技术领域。"
    "本发明首先构建策略图作为中间表示，显式建模代码的模块边界、分支逻辑、状态转换和输入/输出约束，"
    "并通过多轮优化增强策略的完整性和可审计性；然后通过增强测试集和二次验证机制确保语义一致性；"
    "最后利用执行引导反馈和非对称多智能体辩论机制，结合少量高质量示例锚定，深度修复和优化生成的代码，"
    "形成策略构建、代码生成和行为验证相互支持的迭代闭环。本发明能够生成结构清晰、异常处理一致、"
    "易于后续扩展的代码，为大型语言模型在信息系统工程中的应用提供有力支持。"
)
add_paragraph_styled(doc, abstract_text)

# 关键词
kw_p = doc.add_paragraph()
kw_run = kw_p.add_run('关键词：多智能体系统；代码生成；非对称辩论机制；策略建模；软件鲁棒性')
kw_run.font.size = Pt(12)
kw_run.font.name = '宋体'
kw_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
kw_run.bold = True

doc.add_page_break()

# ========== 权利要求书 ==========
add_heading_styled(doc, '权利要求书', level=1)

claims_section = doc.add_paragraph()
claims_section.alignment = WD_ALIGN_PARAGRAPH.CENTER
claims_run = claims_section.add_run('权利要求书')
claims_run.font.size = Pt(16)
claims_run.bold = True
claims_run.font.name = '黑体'
claims_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

claims = [
    "1. 一种基于非对称多智能体辩论的可靠可维护代码生成方法，其特征在于，包括以下步骤：\n"
    "步骤S1：构建策略图作为中间表示，显式建模代码的模块边界、分支逻辑、状态转换和输入/输出约束；\n"
    "步骤S2：将策略图转换为具体代码，并通过增强测试集和二次验证机制确保语义一致性和逻辑正确性；\n"
    "步骤S3：利用执行引导反馈和多智能体辩论机制，结合少量高质量示例锚定，深度修复和优化生成的代码。",
    
    "2. 根据权利要求1所述的方法，其特征在于，步骤S1中所述策略图以有向图形式表示，包含元数据、节点集和边集，"
    "每个节点记录角色、意图、类型化的输入/输出规格以及前置条件和后置条件；每条边指定依赖方向以及可选的守卫条件和优先级。",
    
    "3. 根据权利要求1所述的方法，其特征在于，步骤S1中所述策略图的结构约束包括：图在顶层必须无环；"
    "入口函数后的第一步必须是守卫节点，验证输入并将无效情况路由到显式终止步骤；条件节点限制为二元分支；"
    "非条件节点限制为单条出边；策略图必须完整连通且无不可达节点。",
    
    "4. 根据权利要求1所述的方法，其特征在于，步骤S2中所述增强测试集在基础测试基础上扩展，覆盖边界情况、极端输入和异常状态；"
    "二次验证机制通过CaseCheckAgent验证生成测试用例的逻辑正确性和预期输出准确性。",
    
    "5. 根据权利要求1所述的方法，其特征在于，步骤S3中所述多智能体辩论机制包括：策略专家生成错误-证据-解决方案三元组；"
    "天使智能体支持现有方案，魔鬼智能体提出替代方案进行对抗性讨论；仲裁智能体基于共识分数和方案比较做出最终决策；"
    "风险审计智能体分析潜在副作用和实施挑战。",
    
    "6. 根据权利要求5所述的方法，其特征在于，所述共识分数通过计算智能体间错误类型相似度、证据文本相似度和解决方案文本相似度得到。",
    
    "7. 根据权利要求1所述的方法，其特征在于，步骤S3中所述执行引导反馈包括自动生成并执行测试用例，"
    "将稳定的失败信号转化为对策略图和实现代码的针对性调整，形成策略构建、代码生成和行为验证相互支持的迭代过程。",
    
    "8. 一种基于非对称多智能体辩论的可靠可维护代码生成系统，其特征在于，包括：\n"
    "策略图构建模块，用于构建和优化策略图；\n"
    "代码生成模块，用于将策略图转换为可执行代码；\n"
    "测试增强模块，用于生成和验证增强测试用例；\n"
    "辩论优化模块，用于通过多智能体辩论对代码进行迭代优化。",
    
    "9. 根据权利要求8所述的系统，其特征在于，所述辩论优化模块包括：\n"
    "策略专家单元，用于生成错误-证据-解决方案三元组；\n"
    "天使智能体单元，用于支持现有解决方案；\n"
    "魔鬼智能体单元，用于提出替代方案进行对抗性讨论；\n"
    "仲裁智能体单元，用于基于共识分数和方案比较做出最终决策；\n"
    "风险审计智能体单元，用于分析潜在副作用和实施挑战。",
    
    "10. 一种计算机可读存储介质，其上存储有计算机程序，其特征在于，该程序被处理器执行时实现如权利要求1至7任一项所述方法的步骤。"
]

for idx, claim in enumerate(claims):
    p = doc.add_paragraph()
    run = p.add_run(claim)
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(22)

doc.add_page_break()

# ========== 说明书 ==========
add_heading_styled(doc, '说明书', level=1)

# --- 技术领域 ---
add_heading_styled(doc, '技术领域', level=2)
add_paragraph_styled(doc, 
    "本发明涉及人工智能与代码生成技术领域，具体涉及一种基于非对称多智能体辩论的可靠可维护代码生成方法及系统，"
    "适用于信息系统工程中复杂业务规则的自动化代码生成与优化。")

# --- 背景技术 ---
add_heading_styled(doc, '背景技术', level=2)
bg_text = (
    "大型语言模型在程序合成方面取得了显著进展，展示了在代码补全、程序生成和自动调试方面的强大能力。"
    "然而，在信息系统工程场景中，企业应用通常嵌入复杂的业务规则，涉及多源异构数据、长生命周期状态管理以及严格的安全和合规要求。"
    "代码不仅需要满足功能正确性，还需要在鲁棒性、可维护性和可审计性方面达到高标准。"
    "现有方法存在以下不足：\n\n"
    "（1）缺乏明确的策略层表示，难以系统性描述控制结构、状态演化和输入/输出约束，生成过程中缺乏可分析的中间产物；\n"
    "（2）基础测试集无法促进代码向鲁棒性方向演化，大模型生成的测试用例错误率高；\n"
    "（3）现有多智能体和自修复机制通常在候选代码层面进行局部修补，缺乏对不合理全局决策路径和分支结构的系统性重构能力，"
    "难以确保复杂信息系统逻辑的整体鲁棒性。"
)
add_paragraph_styled(doc, bg_text)

# --- 发明内容 ---
add_heading_styled(doc, '发明内容', level=2)

add_paragraph_styled(doc, "本发明旨在解决现有技术中代码生成结构扁平、分支逻辑不一致、状态管理不稳定等核心问题。", bold=True)

add_paragraph_styled(doc, "为实现上述目的，本发明提供一种基于非对称多智能体辩论的可靠可维护代码生成方法，包括以下步骤：")

steps = [
    "步骤S1：结构去噪层——构建策略图作为中间表示，显式建模代码的模块边界、分支逻辑、状态转换和输入/输出约束；",
    "步骤S2：语义去噪层——将策略图转换为具体代码，并通过增强测试集和二次验证机制确保语义一致性和逻辑正确性；",
    "步骤S3：行为去噪层——利用执行引导反馈和多智能体辩论机制，结合少量高质量示例锚定，深度修复和优化生成的代码。"
]
for step in steps:
    add_paragraph_styled(doc, step)

add_paragraph_styled(doc, 
    "进一步地，步骤S1中策略图以有向图形式表示，包含元数据、节点集和边集，每个节点记录角色、意图、"
    "类型化的输入/输出规格以及前置条件和后置条件；每条边指定依赖方向以及可选的守卫条件和优先级。")

add_paragraph_styled(doc, 
    "进一步地，步骤S2中增强测试集在基础测试基础上扩展，覆盖边界情况、极端输入和异常状态；"
    "二次验证机制通过CaseCheckAgent验证生成测试用例的逻辑正确性和预期输出准确性。")

add_paragraph_styled(doc, 
    "进一步地，步骤S3中非对称辩论过程包括：策略专家生成错误-证据-解决方案三元组；"
    "天使智能体支持现有方案，魔鬼智能体提出替代方案进行对抗性讨论；"
    "仲裁智能体基于共识分数和方案比较做出最终决策；风险审计智能体分析潜在副作用和实施挑战。")

add_paragraph_styled(doc, "对应地，本发明还提供一种基于非对称多智能体辩论的可靠可维护代码生成系统，包括：", bold=True)
system_components = [
    "策略图构建模块，用于构建和优化策略图；",
    "代码生成模块，用于将策略图转换为可执行代码；",
    "测试增强模块，用于生成和验证增强测试用例；",
    "辩论优化模块，用于通过多智能体辩论对代码进行迭代优化。"
]
for comp in system_components:
    add_bullet(doc, comp)

add_paragraph_styled(doc, "本发明具有以下有益效果：", bold=True)
benefits = [
    "（1）通过结构化策略图表示，确保代码结构清晰可分析，在生成前识别并消除错误逻辑和冗余操作，增强了代码的可维护性；",
    "（2）通过增强测试集和执行引导反馈，实现语义去噪，有效评估代码的边界处理能力和运行时问题，确保代码的功能完整性和鲁棒性；",
    "（3）通过非对称多智能体辩论机制，围绕代码质量标准进行批判性评估，识别深层结构问题，通过多轮辩论优化代码决策，确保全局一致性和逻辑完整性。"
]
for b in benefits:
    add_paragraph_styled(doc, b)

# --- 附图说明 ---
add_heading_styled(doc, '附图说明', level=2)
figures = [
    "图1：本发明DeepDebateCoder框架总体架构图；",
    "图2：本发明三层迭代去噪过程示意图；",
    "图3：本发明非对称辩论工作流示意图。"
]
for f in figures:
    add_paragraph_styled(doc, f)

# Try to insert images
fig_paths = [
    OUTPUT_DIR / "Algorithm schematic diagram.png",
    OUTPUT_DIR / "Denoising process.png",
    OUTPUT_DIR / "Debate process.png"
]
fig_names = ["DeepDebateCoder框架总体架构图", "三层迭代去噪过程示意图", "非对称辩论工作流示意图"]

for i, (fp, name) in enumerate(zip(fig_paths, fig_names)):
    if fp.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(fp), width=Inches(5.5))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(f"图{i+1}：{name}")
        cap_run.font.size = Pt(10)
        cap_run.font.name = '宋体'
        cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# --- 具体实施方式 ---
add_heading_styled(doc, '具体实施方式', level=2)

add_heading_styled(doc, '实施例1：DeepDebateCoder框架整体流程', level=3)
emb1 = (
    "如图1所示，本发明框架由三个紧密耦合的层次构成：结构去噪层、语义去噪层和行为去噪层。"
    "整体转换过程可表示为：最终代码 = 合并(行为(语义(结构(问题描述))))。\n\n"
    "结构去噪层：智能体首先推导粗粒度的程序流骨架，捕获解决方案的宏观执行顺序，"
    "如\"输入解析→核心计算→输出格式化\"。基于问题描述和流骨架，智能体以模式约束的JSON格式构建策略图，"
    "包含元数据、节点集和边集。每个节点记录角色、意图、类型化的输入/输出规格以及前置条件和后置条件；"
    "每条边指定依赖方向以及可选的守卫条件和优先级。\n\n"
    "通过执行结构约束进行结构去噪，包括：图在顶层必须无环；入口函数后的第一步必须是守卫节点，"
    "验证输入并将无效情况路由到显式终止步骤；条件节点限制为二元分支，其出边必须携带显式的守卫表达式；"
    "非条件节点限制为单条出边；策略图必须完整连通。\n\n"
    "语义去噪层：将策略图转换为高级元策略表示，确保策略图与元策略之间的严格对齐。"
    "引入增强测试集的二次验证机制，增强测试集在基础测试基础上扩展，覆盖更多边界情况、极端输入和异常状态。"
    "通过CaseCheckAgent验证生成的测试用例，确保每个测试用例的逻辑正确性和预期输出准确性。\n\n"
    "行为去噪层：非对称多智能体辩论机制用于系统性地识别和修复生成代码中的错误。"
    "过程从策略专家开始，为每个失败的测试用例生成由错误-证据-解决方案三元组组成的结构化响应。"
    "基于文本相似度去除重复的错误描述，仅保留唯一问题。"
    "然后，天使智能体和魔鬼智能体基于高质量少量示例对每个问题进行批判性讨论。"
    "天使智能体支持现有解决方案，魔鬼智能体通过提供替代方案进行挑战。"
    "仲裁智能体审查解决方案，利用共识分数指导最终决策。"
    "一旦达成共识，进行最终风险评估，策略专家和代码专家分析潜在副作用和实施挑战。\n\n"
    "实验结果表明，在HumanEval-ET和APPS基准测试上，DeepDebateCoder相对于多个强基线方法表现出稳定优势，"
    "在多个任务上显著提升了pass@1指标。在APPS数据集上，DDCoder在DeepSeek-Chat后端上将pass@1从30-32%提升至66.00%，"
    "在Qwen-Plus后端上从25-29%提升至45.34%。消融实验表明，辩论层对整体性能贡献最大，移除后APPS pass@1下降约12.67个百分点。"
)
add_paragraph_styled(doc, emb1)

add_heading_styled(doc, '实施例2：策略图详细构建方法', level=3)
emb2 = (
    "策略图使用模式约束的JSON格式构建，包含以下字段：\n"
    "metadata：包含问题描述、输入输出规格和全局约束信息；\n"
    "nodes：每个节点包含id、role（如guard、condition、exec、transform、termination）、intent、input_spec、output_spec、preconditions和postconditions；\n"
    "edges：每条边包含source、target、guard_condition（可选）和priority。\n\n"
    "结构约束包括：\n"
    "（1）顶层必须无环，确保全局执行顺序明确；\n"
    "（2）迭代逻辑封装在exec/transform节点内，通过后置条件中的显式退出条件约束；\n"
    "（3）入口函数后的第一步必须是守卫节点；\n"
    "（4）最后一步必须格式化或产生最终输出；\n"
    "（5）条件节点限制为二元分支，出边必须携带互补的守卫表达式；\n"
    "（6）非条件节点限制为单条出边；\n"
    "（7）策略图必须完整连通，无不可达节点、无后继的非终止节点，每个节点产生的值必须被下游节点消费。\n\n"
    "当生成的图形不满足这些要求时，框架拒绝无效结构并请求在同一模式和约束下重新生成，直至约束满足或达到预设的重新生成预算。"
)
add_paragraph_styled(doc, emb2)

add_heading_styled(doc, '实施例3：非对称辩论详细流程', level=3)
emb3 = (
    "如图3所示，非对称辩论工作流始于策略专家A_strategy，为每个失败的测试用例生成错误-证据-解决方案三元组R_i。"
    "基于文本相似度去除重复错误描述，相似度阈值设为0.85。\n\n"
    "然后天使智能体A_angel和魔鬼智能体A_devil基于高质量少量示例对每个问题进行批判性讨论。"
    "天使智能体支持现有解决方案，魔鬼智能体通过提供替代方案进行挑战。"
    "对抗性方法鼓励深度推理，揭示从单一角度可能不明显的弱点。\n"
    "计算智能体间的共识分数，包括错误类型相似度、证据文本相似度和解决方案文本相似度。\n"
    "仲裁智能体A_arbiter基于解决方案比较和共识分数选择最合适的解决方案。\n"
    "最终风险评估由策略专家和代码专家共同完成，确保最终解决方案避免引入新问题或复杂性。"
    "仲裁决策产生可操作的解决方案步骤和相应的风险预防措施。\n\n"
    "在实验中，增强测试用例数量n_test设为10，策略有效性阈值s设为1，策略演化迭代次数t设为3，辩论轮数r设为2。"
    "这些超参数共同决定了模型的优化深度和计算效率。"
)
add_paragraph_styled(doc, emb3)

# --- 实验结果表格 ---
add_heading_styled(doc, '实验结果', level=3)
add_paragraph_styled(doc, "表1展示了DeepDebateCoder在多个基准测试上与基线方法的性能对比（pass@1%）：")

# Create table
table = doc.add_table(rows=10, cols=4, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['方法', 'HumanEval', 'HumanEval-ET', 'APPS']
data = [
    ['Direct', '73.17', '75.00', '32.67'],
    ['CoT', '91.46', '79.90', '18.07'],
    ['Self-Planning', '91.46', '84.10', '28.67'],
    ['Analogical', '92.07', '80.49', '28.00'],
    ['MapCoder', '95.73', '84.10', '32.00'],
    ['CodeSIM', '97.60', '87.20', '30.67'],
    ['DDCoder (Ours)', '92.68', '93.29', '66.00'],
]

# Add headers
for j, h in enumerate(headers):
    cell = table.cell(0, j)
    cell.text = h
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)
    set_cell_shading(cell, 'D9E2F3')

# Add data - DeepSeek-Chat
for i, row_data in enumerate(data):
    for j, val in enumerate(row_data):
        cell = table.cell(i+1, j)
        cell.text = val
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(10)

cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap_run = cap.add_run('表1：DeepSeek-Chat后端pass@1性能对比（%）')
cap_run.font.size = Pt(10)
cap_run.font.name = '宋体'
cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
cap_run.bold = True

# --- 消融实验表格 ---
add_paragraph_styled(doc, "")
add_paragraph_styled(doc, "表2展示了消融实验结果：")

table2 = doc.add_table(rows=6, cols=3, style='Light Grid Accent 1')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = ['变体', 'Pass@1 (%)', '下降 (pp)']
ablation_data = [
    ['w/o 策略图', '64.00', '2.00'],
    ['w/o 增强测试', '64.67', '1.33'],
    ['w/o 辩论', '53.33', '12.67'],
    ['w/o 示例锚定', '63.33', '2.67'],
    ['完整 DDCoder', '66.00', '--'],
]

for j, h in enumerate(headers2):
    cell = table2.cell(0, j)
    cell.text = h
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)
    set_cell_shading(cell, 'D9E2F3')

for i, row_data in enumerate(ablation_data):
    for j, val in enumerate(row_data):
        cell = table2.cell(i+1, j)
        cell.text = val
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(10)

cap2 = doc.add_paragraph()
cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap2_run = cap2.add_run('表2：APPS消融实验结果（pass@1 %）')
cap2_run.font.size = Pt(10)
cap2_run.font.name = '宋体'
cap2_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
cap2_run.bold = True

# ========== 序列表（空） ==========
doc.add_page_break()
add_heading_styled(doc, '序列表', level=1)
add_paragraph_styled(doc, "（无）")

# ========== 保存 ==========
output_path = OUTPUT_DIR / "patent_deepdebatecoder.docx"
doc.save(str(output_path))

# --- Write metadata ---
meta = {
    "run_id": RUN_ID,
    "patent_title": "基于非对称多智能体辩论的可靠可维护代码生成方法及系统",
    "claims_count": 10,
    "figures": [
        "图1: DeepDebateCoder框架总体架构图",
        "图2: 三层迭代去噪过程示意图",
        "图3: 非对称辩论工作流示意图"
    ],
    "output_file": str(output_path.relative_to(SKILL_DIR)),
}
meta_path = OUTPUT_DIR / "patent_meta.json"
meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")

print(f"Patent document generated: {output_path}")
print(f"Metadata written: {meta_path}")
