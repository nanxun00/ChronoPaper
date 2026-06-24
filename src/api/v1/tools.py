import os
import uuid
from pathlib import Path
from fastapi import APIRouter, File, UploadFile, Body, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
from src.utils import setup_logger
from src.utils.paths import UPLOADS_DIR
import shutil

tool = APIRouter(prefix="/tool")
router = tool

logger = setup_logger("server-tools")

DOCUMENTS_DIR = UPLOADS_DIR / "documents"

def ensure_documents_dir():
    DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
    return DOCUMENTS_DIR

class Tool(BaseModel):
    name: str
    title: str
    description: str
    url: str
    method: str

class DocumentGenerateRequest(BaseModel):
    title: str
    content: str
    format: str  # docx, pdf
    template: str = "academic"
    font_family: str = "Times New Roman"
    font_size: int = 12
    line_spacing: float = 1.5
    include_toc: bool = False
    include_page_numbers: bool = True

@tool.get("/", response_model=List[Tool])
async def route_index():
    '''
        返回一个工具列表，共前端页面中展示
        创建一个工具列表，包含两个tool实例
        第一个工具用于文本分块，第二个工具用于pdf转文本

    '''
    tools = [
        Tool(
            name="text-chunking",
            title="文本分块",
            description="将文本分块以更好地理解。可以输入文本或者上传文件。",
            url="/tools/text-chunking",
            method="POST",
        ),
        Tool(
            name="pdf2txt",
            title="PDF转文本",
            description="将PDF文件转换为文本文件。",
            url="/tools/pdf2txt",
            method="POST",
        )
    ]

    return tools

def _generate_docx(body: DocumentGenerateRequest, file_path: Path):
    """生成 DOCX 文档"""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = body.font_family
    font.size = Pt(body.font_size)
    
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(body.title)
    title_run.bold = True
    title_run.font.size = Pt(body.font_size + 4)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    paragraphs = body.content.split('\n')
    for para_text in paragraphs:
        para = doc.add_paragraph(para_text)
        para.paragraph_format.line_spacing = body.line_spacing
        para.paragraph_format.first_line_indent = Inches(0.3)
    
    doc.save(file_path)

def _generate_pdf(body: DocumentGenerateRequest, file_path: Path):
    """生成 PDF 文档"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    
    doc = SimpleDocTemplate(
        str(file_path),
        pagesize=A4,
        rightMargin=25*mm,
        leftMargin=25*mm,
        topMargin=25*mm,
        bottomMargin=25*mm
    )
    
    story = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=body.font_size + 4,
        alignment=TA_CENTER,
        spaceAfter=12*mm,
        fontName='Helvetica-Bold'
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=body.font_size,
        alignment=TA_JUSTIFY,
        fontName='Helvetica',
        spaceAfter=6*mm,
        leading=body.line_spacing * body.font_size
    )
    
    story.append(Paragraph(body.title, title_style))
    story.append(Spacer(1, 6*mm))
    
    paragraphs = body.content.split('\n')
    for para_text in paragraphs:
        if para_text.strip():
            story.append(Paragraph(para_text.strip(), body_style))
    
    doc.build(story)

@tool.post("/generate-document")
async def generate_document(body: DocumentGenerateRequest):
    """生成 DOCX 或 PDF 文档并返回下载链接"""
    if body.format.lower() not in ['docx', 'pdf']:
        raise HTTPException(status_code=400, detail="仅支持 docx 和 pdf 格式")
        
    ensure_documents_dir()
    
    doc_id = uuid.uuid4().hex[:12]
    file_ext = body.format.lower()
    filename = f"{doc_id}.{file_ext}"
    file_path = DOCUMENTS_DIR / filename
    
    logger.info(f"开始生成文档: {filename} (路径: {file_path})")
    
    try:
        if file_ext == 'docx':
            _generate_docx(body, file_path)
        else:
            _generate_pdf(body, file_path)
        
        # 验证文件确实存在
        if not file_path.exists():
            raise RuntimeError(f"文件生成失败: {file_path}")
        
        file_size = os.path.getsize(file_path)
        logger.info(f"文档生成成功: {filename}, 大小: {file_size} bytes")
        
        # 返回完整的 filename 用于下载
        return {
            "doc_id": doc_id,
            "title": body.title,
            "format": body.format,
            "filename": filename,
            "download_url": f"/tool/download-document/{filename}",
            "file_size": file_size,
            "status": "generated"
        }
    except ImportError as e:
        logger.error(f"文档生成依赖未安装: {str(e)}")
        if 'docx' in str(e).lower():
            raise HTTPException(status_code=500, detail="DOCX 生成依赖未安装")
        else:
            raise HTTPException(status_code=500, detail="PDF 生成依赖未安装")
    except Exception as e:
        logger.error(f"文档生成失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"文档生成失败: {str(e)}")

@tool.get("/download-document/{filename}")
async def download_document(filename: str):
    """下载生成的文档"""
    ensure_documents_dir()
    
    # 安全校验：只允许下载 documents 目录下的 .docx 或 .pdf 文件
    if '.' not in filename:
        raise HTTPException(status_code=400, detail="无效的文件名")
    
    file_ext = Path(filename).suffix.lower()
    if file_ext not in ['.docx', '.pdf']:
        raise HTTPException(status_code=400, detail="不支持的文件格式")
    
    file_path = DOCUMENTS_DIR / filename
    
    # 安全检查：防止路径穿越
    try:
        file_path = file_path.resolve()
        if not str(file_path).startswith(str(DOCUMENTS_DIR.resolve())):
            raise HTTPException(status_code=403, detail="禁止访问")
    except HTTPException:
        raise
    
    if not file_path.exists():
        logger.warning(f"下载请求但文件不存在: {filename}")
        raise HTTPException(status_code=404, detail="文档不存在")
    
    media_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' if file_ext == '.docx' else 'application/pdf'
    
    logger.info(f"下载文档: {filename} ({file_path})")
    
    return FileResponse(
        file_path,
        media_type=media_type,
        filename=filename
    )

@tool.post("/text-chunking")
async def text_chunking(text: str = Body(...), params: Dict[str, Any] = Body(...)):
    '''
        文本分块功能的异步API接口
        该函数通过接收post请求，对传入的文本进行分块处理。它使用了外部的chunk函数来进行实际的文本分块操作
        参数：
            - text: 待分块的文本，类型为字符串，通过请求体传递
            - params: 分块操作的参数，类型为字典，包含分块操作所需的额外信息，通过请求体传递
        返回：
        返回一个json对象，包含分块后的文本节点列表，每个节点以字典形式表示
    '''
    from src.services.rag.indexing import chunk
    nodes = chunk(text, params=params)
    return {"nodes": [node.to_dict() for node in nodes]}


@tool.post("/pdf2txt")
async def handle_pdf2txt(file: str = Body(...)):
    '''
        将pdf文件转换为文本
        此函数通过接收一个pdf文件作为输入，然后使用pdf2txt插件将其转换为文本格式
        转换后的文本以json格式返回，包含转换后的文本数据

        参数：
            - file(str)；以字符串形式接受的pdf文件内容
        返回：
            - dict: 包含转换后的文本数据的字典
    '''
    from src.parsers import pdf_simple as pdf2txt
    text = pdf2txt(file, return_text=True)
    return {"text": text}
