"""将文本内容生成为 DOCX / PDF 文件。"""
from __future__ import annotations

import uuid
from pathlib import Path

from pydantic import BaseModel

from src.utils.paths import UPLOADS_DIR

DOCUMENTS_DIR = UPLOADS_DIR / "documents"


class DocumentGenerateRequest(BaseModel):
    title: str
    content: str
    format: str = "docx"
    template: str = "academic"
    font_family: str = "Times New Roman"
    font_size: int = 12
    line_spacing: float = 1.5
    include_toc: bool = False
    include_page_numbers: bool = True


def ensure_documents_dir() -> Path:
    DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
    return DOCUMENTS_DIR


def _generate_docx(body: DocumentGenerateRequest, file_path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = body.font_family
    style.font.size = Pt(body.font_size)

    title_para = doc.add_paragraph()
    title_run = title_para.add_run(body.title)
    title_run.bold = True
    title_run.font.size = Pt(body.font_size + 4)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for para_text in body.content.split("\n"):
        if not para_text.strip():
            continue
        para = doc.add_paragraph(para_text)
        para.paragraph_format.line_spacing = body.line_spacing
        para.paragraph_format.first_line_indent = Inches(0.3)

    doc.save(file_path)


def _generate_pdf(body: DocumentGenerateRequest, file_path: Path) -> None:
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    doc = SimpleDocTemplate(
        str(file_path),
        pagesize=A4,
        rightMargin=25 * mm,
        leftMargin=25 * mm,
        topMargin=25 * mm,
        bottomMargin=25 * mm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Heading1"],
        fontSize=body.font_size + 4,
        alignment=TA_CENTER,
        spaceAfter=12 * mm,
        fontName="Helvetica-Bold",
    )
    body_style = ParagraphStyle(
        "CustomBody",
        parent=styles["Normal"],
        fontSize=body.font_size,
        alignment=TA_JUSTIFY,
        fontName="Helvetica",
        spaceAfter=6 * mm,
        leading=body.line_spacing * body.font_size,
    )

    story = [Paragraph(body.title, title_style), Spacer(1, 6 * mm)]
    for para_text in body.content.split("\n"):
        if para_text.strip():
            story.append(Paragraph(para_text.strip(), body_style))
    doc.build(story)


def generate_document_file(
    *,
    title: str,
    content: str,
    format: str = "docx",
    font_family: str = "Times New Roman",
    font_size: int = 12,
    line_spacing: float = 1.5,
) -> dict:
    """生成文档并返回下载元数据。"""
    fmt = (format or "docx").lower()
    if fmt not in ("docx", "pdf"):
        raise ValueError("仅支持 docx 和 pdf 格式")

    ensure_documents_dir()
    doc_id = uuid.uuid4().hex[:12]
    filename = f"{doc_id}.{fmt}"
    file_path = DOCUMENTS_DIR / filename

    body = DocumentGenerateRequest(
        title=title.strip() or "Untitled",
        content=content.strip(),
        format=fmt,
        font_family=font_family,
        font_size=font_size,
        line_spacing=line_spacing,
    )

    if fmt == "docx":
        _generate_docx(body, file_path)
    else:
        _generate_pdf(body, file_path)

    if not file_path.exists():
        raise RuntimeError(f"文件生成失败: {file_path}")

    safe_title = "".join(c for c in body.title if c not in '\\/:*?"<>|').strip() or "document"
    display_name = f"{safe_title}.{fmt}"
    file_size = file_path.stat().st_size

    return {
        "doc_id": doc_id,
        "title": body.title,
        "format": fmt,
        "filename": filename,
        "name": display_name,
        "download_url": f"/tool/download-document/{filename}",
        "url": f"/tool/download-document/{filename}",
        "kind": "file",
        "file_size": file_size,
        "status": "generated",
    }
