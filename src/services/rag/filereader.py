import os
import re

from pathlib import Path

import markdown
import win32com
from docx import Document
from langchain_community.document_loaders import (
    TextLoader,
    Docx2txtLoader,
    UnstructuredMarkdownLoader,
    UnstructuredHTMLLoader, UnstructuredWordDocumentLoader,
)

from langchain.text_splitter import RecursiveCharacterTextSplitter

def pdfreader(file_path):
    """读取PDF文件并返回text文本"""
    assert os.path.exists(file_path), "File not found"
    assert file_path.endswith(".pdf"), "File format not supported"

    from llama_index.readers.file import PDFReader
    doc = PDFReader().load_data(file=Path(file_path))

    # 简单的拼接起来之后返回纯文本
    text = "\n\n".join([d.get_content() for d in doc])
    print("doc: {}", text)
    return text

def plainreader_new(file_path):
    """
        使用文件解析器将文件切分成固定大小的块
        Args:
            file_path: 文件路径
        """
    params = None
    chunk_size = int(params.get("chunk_size", 500))
    chunk_overlap = int(params.get("chunk_overlap", 100))

    file_type = Path(file_path).suffix.lower()
    # 选择合适的加载器
    if file_type in ['.txt']:
        loader = TextLoader(file_path)

    elif file_type in ['.md']:
        loader = UnstructuredMarkdownLoader(file_path)

    elif file_type in ['.docx', '.doc']:
        loader = Docx2txtLoader(file_path)

    elif file_type in ['.html', '.htm']:
        loader = UnstructuredHTMLLoader(file_path)

    else:
        raise ValueError(f"不支持的文件类型: {file_type}")

    # 加载文档
    docs = loader.load()

    # 创建文本分割器
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ".", " ", ""],
    )
    # 分割文档
    nodes = text_splitter.split_documents(docs)
    return nodes


## 已优化
def plainreader(file_path):
    """读取普通文本文件并返回text文本"""
    # assert os.path.exists(file_path), "File not found"
    #
    # with open(file_path, "r") as f:
    #     text = f.read()
    print("66666666")
    """读取普通文本文件并返回text文本"""
    assert os.path.exists(file_path), "File not found"
    file_type = Path(file_path).suffix.lower()
    print(f"文件类型：77777， {file_type}")
    # 使用LangChain的文本加载器
    # 选择合适的加载器
    text =""
    if file_type in ['.txt']:
        loader = TextLoader(file_path)
        docs = loader.load()
        text = "\n\n".join([d.page_content for d in docs])

    elif file_type in ['.md']:
        print("9999999")
        # md文档存在问题
        with open(file_path, "r", encoding="utf-8") as md_file:
            markdown_text = md_file.read()
        converted_text = markdown.markdown(markdown_text)
        text = re.sub('<[^>]+>', '', converted_text)

    # elif file_type in ['.doc']:
    #     print("888888")
    #     doc = Document(file_path)
    #     text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs])

    elif file_type in ['.docx']:
        print("888888")
        try:
            # 尝试使用python-docx打开文件
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            print(f"成功加载: {file_path}")
        except Exception as e:
            print(f"无法加载文件 {file_path}: {e}")
            print("bbb")
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(file_path)
            text = doc.Content.Text
            doc.Close(False)
            word.Quit()


        # try:
        #     print("第二次")
        #     # 尝试使用python-docx打开文件
        #     doc = Document(file_path)
        #     text = "\n".join([para.text for para in doc.paragraphs])
        #
        #     print(f"成功加载: {file_path}")
        # except Exception as e:
        #     print(f"无法加载文件 {file_path}: {e}")
        # doc = Document(file_path)
        # text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs])
        print(text)

    #
    # elif file_type in ['.docx']:
    #     print("888888")
    #     # file_path zip文件
    #     loader = UnstructuredWordDocumentLoader(file_path)
    #     # loader = Docx2txtLoader(file_path)
    #     text = loader.load()

    # elif file_type in ['.html', '.htm']:
    #     loader = UnstructuredHTMLLoader(file_path)

    else:
        raise ValueError(f"不支持的文件类型: {file_type}")

    # print("loader",loader)
    # 这里未进行文本分割，直接进行了拼接

    print(f"doc: {text}")
    return text

if __name__ == "__main__":
    from docx import Document
    import os
    a = "src/saves/pdf2txt/"
    path = "src/saves/pdf2txt/1_9f215ec030410c08.docx"
    new_path = path.split("src/saves/pdf2txt/")
    b = a+"new"+new_path[1]
    print(b)


